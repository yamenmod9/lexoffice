from __future__ import annotations

import base64
from collections import defaultdict
from datetime import date
from io import BytesIO

from docx import Document as WordDocument
from flask import Blueprint, Response, g, request
from sqlalchemy import func

from app.models import Case, Client, HearingSession, Invoice, Judgment, Payment, Task
from app.services.pdf_service import build_simple_pdf
from app.utils.decorators import require_permission
from app.utils.responses import fail, ok
from app.utils.serialization import model_to_dict

bp = Blueprint("reports", __name__, url_prefix="/api/v1/reports")


def _forbid_financial_scope_for_operational_reports():
    if getattr(g, "permission_scope", "all") == "financial":
        return fail("FORBIDDEN", "Insufficient permissions", status=403)
    return None


@bp.get("/financial-monthly")
@require_permission("reports", "full")
def financial_monthly_pdf():
    year = int(request.args.get("year", date.today().year))
    month = int(request.args.get("month", date.today().month))

    start = date(year, month, 1)
    end = date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)

    payments = (
        Payment.query.filter_by(office_id=g.current_user.office_id)
        .filter(Payment.payment_date >= start, Payment.payment_date < end)
        .all()
    )
    invoices = (
        Invoice.query.filter_by(office_id=g.current_user.office_id)
        .filter(Invoice.issue_date >= start, Invoice.issue_date < end)
        .all()
    )

    lines = [
        f"Month: {year}-{month:02d}",
        f"Payments count: {len(payments)}",
        f"Invoices count: {len(invoices)}",
        f"Collected total: {sum(item.amount for item in payments)}",
        f"Invoiced total: {sum(item.total for item in invoices)}",
    ]
    pdf = build_simple_pdf("Monthly Financial Report", lines)
    return Response(pdf, mimetype="application/pdf")


@bp.get("/lawyer-performance")
@require_permission("reports", "full")
def lawyer_performance():
    denied = _forbid_financial_scope_for_operational_reports()
    if denied:
        return denied

    office_id = g.current_user.office_id
    cases = Case.query.filter_by(office_id=office_id).all()
    tasks = Task.query.filter_by(office_id=office_id).all()

    result = defaultdict(lambda: {"cases": 0, "tasks": 0, "done_tasks": 0, "deadlines_missed": 0})
    for case in cases:
        key = str(case.responsible_lawyer_id)
        result[key]["cases"] += 1

    for task in tasks:
        key = str(task.assigned_to)
        result[key]["tasks"] += 1
        if str(getattr(task.status, "value", task.status)) == "done":
            result[key]["done_tasks"] += 1
        if task.deadline and str(getattr(task.status, "value", task.status)) != "done" and task.deadline.date() < date.today():
            result[key]["deadlines_missed"] += 1

    return ok(data=result)


@bp.get("/workload")
@require_permission("reports", "full")
def workload():
    denied = _forbid_financial_scope_for_operational_reports()
    if denied:
        return denied

    office_id = g.current_user.office_id
    tasks = Task.query.filter_by(office_id=office_id).all()

    per_user = defaultdict(lambda: {"total": 0, "urgent": 0, "in_progress": 0})
    for task in tasks:
        key = str(task.assigned_to)
        per_user[key]["total"] += 1
        if str(getattr(task.priority, "value", task.priority)) == "urgent":
            per_user[key]["urgent"] += 1
        if str(getattr(task.status, "value", task.status)) == "in_progress":
            per_user[key]["in_progress"] += 1

    return ok(data=per_user)


@bp.get("/win-loss")
@require_permission("reports", "full")
def win_loss():
    denied = _forbid_financial_scope_for_operational_reports()
    if denied:
        return denied

    office_id = g.current_user.office_id
    judgments = Judgment.query.filter_by(office_id=office_id).all()

    by_type = defaultdict(lambda: {"win": 0, "loss": 0, "other": 0})
    by_court = defaultdict(lambda: {"win": 0, "loss": 0, "other": 0})
    by_lawyer = defaultdict(lambda: {"win": 0, "loss": 0, "other": 0})

    for judgment in judgments:
        case = Case.query.filter_by(id=judgment.case_id, office_id=office_id).first()
        result = str(getattr(judgment.result, "value", judgment.result))
        bucket = "win" if result in {"full_win", "partial_win"} else "loss" if result == "loss" else "other"

        if case:
            by_type[str(getattr(case.case_type, "value", case.case_type))][bucket] += 1
            by_lawyer[str(case.responsible_lawyer_id)][bucket] += 1
        by_court[judgment.court or "unknown"][bucket] += 1

    return ok(data={"by_type": by_type, "by_court": by_court, "by_lawyer": by_lawyer})


@bp.get("/client-statement/<uuid:client_id>")
@require_permission("reports", "full")
def client_statement_files(client_id):
    office_id = g.current_user.office_id
    client = Client.query.filter_by(id=client_id, office_id=office_id, is_deleted=False).first()
    if not client:
        return fail("NOT_FOUND", "Client not found", status=404)

    payments = Payment.query.filter_by(office_id=office_id, client_id=client_id).all()
    invoices = Invoice.query.filter_by(office_id=office_id, client_id=client_id).all()

    lines = [
        f"Client: {client.full_name_ar}",
        f"Client Number: {client.client_number}",
        f"Total invoices: {sum(item.total for item in invoices)}",
        f"Total payments: {sum(item.amount for item in payments)}",
        f"Outstanding: {max(sum(item.total for item in invoices) - sum(item.amount for item in payments), 0)}",
    ]
    pdf = build_simple_pdf("Client Statement", lines)

    doc = WordDocument()
    doc.add_heading("Client Statement", level=1)
    for line in lines:
        doc.add_paragraph(line)
    for invoice in invoices:
        doc.add_paragraph(f"Invoice {invoice.invoice_number}: {invoice.total}")
    for payment in payments:
        doc.add_paragraph(f"Payment {payment.payment_date}: {payment.amount}")

    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    return ok(
        data={
            "pdf_base64": base64.b64encode(pdf).decode("utf-8"),
            "docx_base64": base64.b64encode(docx_bytes).decode("utf-8"),
        }
    )


@bp.get("/sessions-report")
@require_permission("reports", "full")
def sessions_report():
    denied = _forbid_financial_scope_for_operational_reports()
    if denied:
        return denied

    office_id = g.current_user.office_id

    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    query = HearingSession.query.filter_by(office_id=office_id)
    if date_from:
        query = query.filter(HearingSession.session_date >= date.fromisoformat(date_from))
    if date_to:
        query = query.filter(HearingSession.session_date <= date.fromisoformat(date_to))

    sessions = query.order_by(HearingSession.session_date.asc()).all()

    summary = {
        "total_sessions": len(sessions),
        "by_result": defaultdict(int),
    }
    for item in sessions:
        summary["by_result"][str(getattr(item.result, "value", item.result or "pending"))] += 1

    return ok(data={"summary": summary, "sessions": [model_to_dict(item) for item in sessions]})
